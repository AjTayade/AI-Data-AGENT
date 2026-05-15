'''import streamlit as st
import pandas as pd
import pathlib
import sqlite3
import tempfile
import PyPDF2
import docx
import re
import numpy as np
import json
import io
import time
import google.generativeai as genai
from supabase import create_client, Client
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_experimental.agents import create_pandas_dataframe_agent

st.set_page_config(page_title="AI Data Agent SaaS", layout="wide")

# ──────────────────────────────────────────────────────────────
# 1. SETUP THE BRAIN & THE VAULT
# ──────────────────────────────────────────────────────────────
try:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    model = genai.GenerativeModel('gemini-2.0-flash')
    supabase: Client = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])

    if 'session' in st.session_state and st.session_state['session'] is not None:
        try:
            supabase.postgrest.auth(st.session_state['session'].access_token)
            supabase.auth.set_session(
                st.session_state['session'].access_token,
                st.session_state['session'].refresh_token
            )
        except Exception:
            st.session_state['user'] = None
            st.session_state['session'] = None
except Exception as e:
    st.error(f"Setup Error: Please check your Streamlit Secrets. ({e})")
    st.stop()

# ──────────────────────────────────────────────────────────────
# 2. THE BOUNCER (LOGIN / REGISTER)
# ──────────────────────────────────────────────────────────────
if 'user' not in st.session_state:
    st.session_state['user'] = None
if 'session' not in st.session_state:
    st.session_state['session'] = None

if st.session_state['user'] is None:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.title("🔐 Welcome to AI Data Agent")
        st.markdown("Please log in to access your secure workspace.")

        auth_tab1, auth_tab2 = st.tabs(["Login", "Register"])
        with auth_tab1:
            log_email = st.text_input("Email", key="log_email")
            log_pass  = st.text_input("Password", type="password", key="log_pass")
            if st.button("Login", use_container_width=True, type="primary"):
                try:
                    response = supabase.auth.sign_in_with_password({"email": log_email, "password": log_pass})
                    st.session_state['user']    = response.user
                    st.session_state['session'] = response.session
                    st.rerun()
                except Exception:
                    st.error("Login failed. Please check your credentials.")

        with auth_tab2:
            reg_name  = st.text_input("First Name", key="reg_name")
            reg_email = st.text_input("New Email", key="reg_email")
            reg_pass  = st.text_input("New Password", type="password", key="reg_pass")
            if st.button("Register", use_container_width=True):
                try:
                    supabase.auth.sign_up({
                        "email": reg_email, "password": reg_pass,
                        "options": {"data": {"first_name": reg_name}}
                    })
                    st.success("Registration successful! You can now log in.")
                except Exception as e:
                    st.error(f"Registration failed: {e}")
    st.stop()

# ──────────────────────────────────────────────────────────────
# 2.5  SESSION RESTORE  (runs once per login after any refresh)
# ──────────────────────────────────────────────────────────────
# On a hard refresh Streamlit wipes st.session_state, so notebook_list
# and active_notebook disappear even though the data is safe in Supabase.
# This block rebuilds them from Supabase the first time the user lands on
# the page after a refresh — before anything in the sidebar renders.
_uid = st.session_state['user'].id

if not st.session_state.get('_session_restored'):
    try:
        # Collect every distinct notebook_name this user owns across both tables
        _raw_nb   = supabase.table("raw_datasets").select("notebook_name") \
                        .eq("user_id", _uid).execute()
        _clean_nb = supabase.table("cleaned_datasets").select("notebook_name") \
                        .eq("user_id", _uid).execute()

        _all_nb = sorted(set(
            [r['notebook_name'] for r in _raw_nb.data]  +
            [r['notebook_name'] for r in _clean_nb.data]
        ))

        # Restore notebook list (don't overwrite if the user already has one
        # in this session, e.g. they just created a new notebook)
        if _all_nb and not st.session_state.get('notebook_list'):
            st.session_state['notebook_list'] = _all_nb

        # Restore active notebook — default to the first one found
        if (st.session_state.get('notebook_list')
                and st.session_state.get('active_notebook') is None):
            st.session_state['active_notebook'] = st.session_state['notebook_list'][0]

        st.session_state['_session_restored'] = True

    except Exception as _restore_err:
        # Non-fatal: user can still create a new notebook manually
        st.session_state['_session_restored'] = True

# ──────────────────────────────────────────────────────────────
# 3. THE SIDEBAR (PROFILE & NOTEBOOKS)
# ──────────────────────────────────────────────────────────────
user_name = st.session_state['user'].user_metadata.get('first_name', 'User')

st.sidebar.markdown("### 👤 My Profile")
st.sidebar.markdown(f"**Name:** {user_name}")
st.sidebar.markdown(f"**Email:** {st.session_state['user'].email}")

if st.sidebar.button("🚪 Log Out", use_container_width=True):
    # Clear everything so the next login starts with a clean restore
    for _k in ['user', 'session', 'notebook_list', 'active_notebook',
               'raw_cloud', 'clean_cloud', 'messages', '_session_restored']:
        st.session_state.pop(_k, None)
    supabase.auth.sign_out()
    st.rerun()

st.sidebar.divider()

if 'notebook_list' not in st.session_state:
    st.session_state['notebook_list'] = []
if 'active_notebook' not in st.session_state:
    st.session_state['active_notebook'] = None

st.sidebar.markdown("### 📚 My Notebooks")

# ── New Notebook Modal ──────────────────────────────────────
if 'show_nb_modal' not in st.session_state:
    st.session_state['show_nb_modal'] = False

if st.sidebar.button("📓 New Notebook", use_container_width=True):
    st.session_state['show_nb_modal'] = True

@st.dialog("Create a New Notebook")
def new_notebook_modal():
    nb_input = st.text_input("Notebook name", placeholder="e.g., Q3 Financials")
    col_ok, col_cancel = st.columns(2)
    with col_ok:
        if st.button("✅ Create", use_container_width=True, type="primary"):
            if nb_input and nb_input not in st.session_state['notebook_list']:
                st.session_state['notebook_list'].append(nb_input)
                st.session_state['active_notebook'] = nb_input
                st.session_state['show_nb_modal']   = False
                st.rerun()
            elif nb_input in st.session_state['notebook_list']:
                st.warning("A notebook with this name already exists.")
            else:
                st.warning("Please enter a name.")
    with col_cancel:
        if st.button("Cancel", use_container_width=True):
            st.session_state['show_nb_modal'] = False
            st.rerun()

if st.session_state['show_nb_modal']:
    new_notebook_modal()

if st.session_state['notebook_list']:
    st.sidebar.divider()
    st.sidebar.markdown("**Switch Notebook:**")
    st.session_state['active_notebook'] = st.sidebar.radio(
        "Select active workspace:",
        st.session_state['notebook_list'],
        label_visibility="collapsed"
    )

# ──────────────────────────────────────────────────────────────
# 4. THE NOTEBOOK ENFORCER
# ──────────────────────────────────────────────────────────────
if st.session_state['active_notebook'] is None:
    st.warning("👋 Welcome! Please create your first Notebook in the sidebar to begin working.")
    st.stop()

st.title(f"📂 Notebook: {st.session_state['active_notebook']}")
st.divider()

# ──────────────────────────────────────────────────────────────
# 5. SECURE CLOUD FETCHING
# ──────────────────────────────────────────────────────────────
user_id = st.session_state['user'].id
nb_name = st.session_state['active_notebook']

# Only re-fetch from Supabase when the active notebook changes OR when
# raw_cloud/clean_cloud are missing (e.g. first load after a refresh).
# This prevents wiping already-loaded data on every Streamlit re-run.
_nb_changed = (st.session_state.get('_loaded_nb') != nb_name)

if _nb_changed or 'raw_cloud' not in st.session_state:
    st.session_state['raw_cloud']   = {}
    st.session_state['clean_cloud'] = {}

    with st.spinner(f"Loading {nb_name} secure vault..."):
        try:
            raw_res = supabase.table("raw_datasets").select("file_name, data") \
                .eq("user_id", user_id).eq("notebook_name", nb_name).execute()
            st.session_state['raw_cloud'] = {
                r['file_name']: pd.DataFrame(r['data']) for r in raw_res.data
            }

            clean_res = supabase.table("cleaned_datasets").select("file_name, data") \
                .eq("user_id", user_id).eq("notebook_name", nb_name).execute()
            st.session_state['clean_cloud'] = {
                r['file_name']: pd.DataFrame(r['data']) for r in clean_res.data
            }

            # Mark which notebook is currently loaded so we don't re-fetch unnecessarily
            st.session_state['_loaded_nb'] = nb_name

        except Exception as e:
            st.warning(f"Could not fetch cloud data: {e}")

# ──────────────────────────────────────────────────────────────
# TABBED ARCHITECTURE
# ──────────────────────────────────────────────────────────────
tab1, tab2, tab3 = st.tabs([
    "📤 Part 1 — Upload",
    "🧹 Part 2 — Sanitise & Query",
    "📊 Part 3 — Dashboard"
])


# ══════════════════════════════════════════════════════════════
# TAB 1 — UPLOAD
# ══════════════════════════════════════════════════════════════
with tab1:
    st.subheader("📤 Upload New Data")
    st.caption(
        "Supported formats: CSV, Excel (.xlsx/.xls), JSON, SQLite (.db/.sqlite), "
        "PDF, Word (.docx), and plain text (.txt). You can upload multiple files at once."
    )

    uploaded_files = st.file_uploader(
        f"Drop files into '{nb_name}'",
        type=["csv", "xlsx", "xls", "json", "db", "sqlite", "pdf", "docx", "txt"],
        accept_multiple_files=True
    )

    if st.button("🚀 Upload File(s) to Database", type="primary"):
        if not uploaded_files:
            st.warning("Please select at least one file to upload.")
        else:
            for file in uploaded_files:

                # ── GUARD 1: Duplicate in this notebook ────────────────
                if file.name in st.session_state.get('raw_cloud', {}):
                    st.warning(
                        f"⚠️ **'{file.name}'** already exists in this notebook. "
                        f"Upload skipped to avoid duplicates. "
                        f"To replace it, delete the existing file using the file manager below."
                    )
                    continue

                # ── GUARD 2: File exists in another notebook ───────────
                try:
                    existing_res = supabase.table("raw_datasets") \
                        .select("file_name, data, notebook_name") \
                        .eq("user_id", user_id) \
                        .eq("file_name", file.name) \
                        .neq("notebook_name", nb_name) \
                        .limit(1) \
                        .execute()

                    if existing_res.data:
                        source_nb  = existing_res.data[0]['notebook_name']
                        reused_df  = pd.DataFrame(existing_res.data[0]['data'])
                        reused_df  = reused_df.replace({np.nan: None})
                        supabase.table('raw_datasets').insert({
                            'file_name':     file.name,
                            'data':          reused_df.to_dict(orient='records'),
                            'user_id':       user_id,
                            'notebook_name': nb_name
                        }).execute()
                        st.session_state['raw_cloud'][file.name] = reused_df
                        st.info(
                            f"♻️ **'{file.name}'** was found in your notebook **'{source_nb}'**. "
                            f"Linked it here — no duplicate data stored."
                        )
                        continue  # skip normal parse+upload for this file
                except Exception as cross_err:
                    st.caption(
                        f"Note: Could not check other notebooks for '{file.name}'. "
                        f"Uploading fresh copy. ({cross_err})"
                    )

                # ── NORMAL PARSE + UPLOAD ──────────────────────────────
                file_ext = pathlib.Path(file.name).suffix.lower()
                try:
                    if file_ext in ['.csv', '.xlsx', '.xls', '.json']:
                        if file_ext == '.csv':
                            df = pd.read_csv(file)
                        elif file_ext in ['.xlsx', '.xls']:
                            df = pd.read_excel(file)
                        elif file_ext == '.json':
                            try:
                                df = pd.read_json(file, orient='records')
                            except Exception:
                                file.seek(0)
                                try:
                                    df = pd.read_json(file, lines=True)
                                except Exception:
                                    file.seek(0)
                                    content = file.getvalue().decode("utf-8")
                                    valid_json_objects = []
                                    for line in content.splitlines():
                                        if line.strip():
                                            try:
                                                valid_json_objects.append(json.loads(line))
                                            except json.JSONDecodeError:
                                                pass
                                    if valid_json_objects:
                                        df = pd.json_normalize(valid_json_objects)
                                    else:
                                        raise Exception("Could not extract tabular JSON data.")

                        df = df.replace({np.nan: None})
                        with st.spinner(f"Saving {file.name}..."):
                            supabase.table('raw_datasets').insert({
                                'file_name':     file.name,
                                'data':          df.to_dict(orient='records'),
                                'user_id':       user_id,
                                'notebook_name': nb_name
                            }).execute()
                        st.session_state['raw_cloud'][file.name] = df
                        st.success(f"✅ Uploaded: {file.name}")

                    elif file_ext in ['.txt', '.pdf', '.docx', '.db', '.sqlite']:
                        if file_ext == '.txt':
                            text_data = file.getvalue().decode("utf-8")
                        elif file_ext == '.pdf':
                            pdf_reader = PyPDF2.PdfReader(file)
                            text_data  = "".join([page.extract_text() + "\n" for page in pdf_reader.pages])
                        elif file_ext == '.docx':
                            doc_obj   = docx.Document(file)
                            text_data = "\n".join([para.text for para in doc_obj.paragraphs])
                        elif file_ext in ['.db', '.sqlite']:
                            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
                                tmp.write(file.read())
                                tmp_path = tmp.name
                            conn   = sqlite3.connect(tmp_path)
                            tables = pd.read_sql("SELECT name FROM sqlite_master WHERE type='table'", conn)
                            text_data = f"SQLite DB: {file.name}\nTables: {tables['name'].tolist()}"
                            conn.close()

                        df = pd.DataFrame([{"document_name": file.name, "content": text_data}])
                        with st.spinner(f"Saving {file.name}..."):
                            supabase.table('raw_datasets').insert({
                                'file_name':     file.name,
                                'data':          df.to_dict(orient='records'),
                                'user_id':       user_id,
                                'notebook_name': nb_name
                            }).execute()
                        st.session_state['raw_cloud'][file.name] = df
                        st.success(f"✅ Uploaded document: {file.name}")

                except Exception as e:
                    st.error(f"Error uploading {file.name}: {e}")

    # ── Success nudge ──────────────────────────────────────────
    if st.session_state.get('raw_cloud'):
        st.success(
            f"✅ {len(st.session_state['raw_cloud'])} file(s) ready — "
            f"head to **Part 2** to sanitise or **Part 3** for your dashboard."
        )

    # ── File Preview ───────────────────────────────────────────
    if st.session_state.get('raw_cloud'):
        st.divider()
        st.subheader("🗃️ Uploaded Files Preview")
        for fname, fdf in st.session_state['raw_cloud'].items():
            with st.expander(f"📄 {fname}  —  {len(fdf):,} rows × {len(fdf.columns)} cols"):
                st.dataframe(fdf.head(5), use_container_width=True)

    # ── File Manager (Delete) ──────────────────────────────────
    st.divider()
    with st.expander("🗂️ Manage & Delete Files", expanded=False):
        all_stored = {
            **{f"[RAW]   {k}":   ("raw_datasets",     k) for k in st.session_state.get('raw_cloud', {}).keys()},
            **{f"[CLEAN] {k}":   ("cleaned_datasets",  k) for k in st.session_state.get('clean_cloud', {}).keys()},
        }

        if not all_stored:
            st.caption("No files in this notebook yet.")
        else:
            st.caption(
                "Select a file and press Delete to permanently remove it "
                "from the database and this notebook."
            )
            file_to_delete = st.selectbox(
                "Select file to delete:",
                list(all_stored.keys()),
                key="delete_selector"
            )

            if st.button("🗑️ Delete Selected File", type="secondary", use_container_width=True):
                table_name, original_name = all_stored[file_to_delete]
                try:
                    supabase.table(table_name).delete() \
                        .eq("user_id",       user_id) \
                        .eq("notebook_name", nb_name) \
                        .eq("file_name",     original_name) \
                        .execute()

                    if table_name == "raw_datasets" and original_name in st.session_state['raw_cloud']:
                        del st.session_state['raw_cloud'][original_name]
                    elif table_name == "cleaned_datasets" and original_name in st.session_state['clean_cloud']:
                        del st.session_state['clean_cloud'][original_name]

                    st.success(f"✅ '{original_name}' permanently deleted.")
                    st.rerun()
                except Exception as e:
                    st.error(f"🚨 Could not delete file: {e}")


# ══════════════════════════════════════════════════════════════
# TAB 2 — SANITISE & CHAT
# ══════════════════════════════════════════════════════════════
with tab2:

    # ── Empty state ────────────────────────────────────────────
    if not st.session_state.get('raw_cloud'):
        st.info("⬆️ No data yet — upload your files in **Part 1** first.")

    # ── Sanitizer ─────────────────────────────────────────────
    st.header("🧠 Universal Data Sanitizer")

    raw_keys = list(st.session_state.get('raw_cloud', {}).keys())
    if raw_keys:
        selected_files = st.multiselect(
            "Select raw dataset(s) to sanitise:",
            raw_keys,
            default=raw_keys[:1]
        )

        if st.button("🪄 Run Exhaustive AI Sanitization", type="primary"):
            if not selected_files:
                st.warning("Please select at least one file to sanitize.")
            else:
                total_files  = len(selected_files)
                progress_bar = st.progress(0)
                status_text  = st.empty()

                for i, file in enumerate(selected_files):
                    est_time = (total_files - i) * 15
                    status_text.info(
                        f"Processing {i+1} of {total_files}: **{file}** "
                        f"(Est. remaining: ~{est_time}s)"
                    )

                    df_to_process  = st.session_state['raw_cloud'][file]
                    original_rows  = len(df_to_process)
                    original_cols  = len(df_to_process.columns)

                    # ── Build rich column profile for the AI ──────────
                    profile_lines = []
                    for col in df_to_process.columns:
                        dtype      = str(df_to_process[col].dtype)
                        null_pct   = round(df_to_process[col].isnull().mean() * 100, 1)
                        sample_vals = (
                            df_to_process[col].dropna().astype(str).unique()[:5].tolist()
                        )
                        profile_lines.append(
                            f"  - {col} | dtype={dtype} | nulls={null_pct}% | samples={sample_vals}"
                        )
                    col_profile = "\n".join(profile_lines)
                    full_sample = df_to_process.head(20).to_csv(index=False)

                    # ── Comprehensive AI sanitization prompt ──────────
                    clean_prompt = f"""
You are a Senior Data Engineer and experienced Data Analyst. Your task is to write a single Python
function called `clean_data(df)` that exhaustively sanitizes the provided dataset for analysis.

DATASET PROFILE (use this to make targeted, column-level decisions):
{col_profile}

SAMPLE DATA (first 20 rows):
{full_sample}

YOUR FUNCTION MUST HANDLE ALL OF THE FOLLOWING — apply only what is relevant to this specific data:

── STRUCTURAL ──
1. Convert all column headers to strict snake_case: lowercase, strip leading/trailing whitespace,
   replace spaces and special characters with underscores, collapse multiple underscores.
2. Drop columns that are 100% empty (all NaN/None).
3. Drop rows that are 100% empty.
4. Drop exact duplicate rows.
5. Reset the DataFrame index after dropping rows.

── WHITESPACE & ENCODING ──
6. For every object/string column: strip leading and trailing whitespace using df[col].str.strip().
7. Remove hidden/invisible characters: null bytes (\\x00), BOM (\\ufeff), zero-width spaces (\\u200b),
   carriage returns (\\r). Use df[col].str.replace() with a compiled regex per column.
8. Normalize multiple consecutive internal spaces to a single space using regex.

── NULL / MISSING VALUES ──
9. Replace all known junk null representations with np.nan:
   Exact matches (case-insensitive): '', 'n/a', 'na', 'none', 'null', 'nil', 'nan', '-', '--',
   '---', '?', 'unknown', 'undefined', 'missing', 'not available', 'not applicable', '#n/a', '#null!'.
10. For numeric columns: if null rate < 20%, impute with column median.
    If null rate is 20–50%, leave as NaN. If null rate > 50%, drop the column and
    print(f"Dropped high-null column: {{col}}").

── NUMERIC COLUMNS ──
11. If a column that should be numeric contains strings with currency symbols ($, €, £, ¥),
    comma thousands separators, or percentage signs (%), strip them and cast to float.
12. If a column has dtype object but all non-null values are numeric strings, cast it to float.
13. Detect and cap outliers using IQR (1.5x rule): replace values below Q1-1.5*IQR
    or above Q3+1.5*IQR with the boundary value (Winsorize — do NOT drop rows).
    Only apply where outlier count > 1% of total rows.

── DATE / TIME COLUMNS ──
14. Detect date columns by name keywords: 'date', 'time', 'dt', 'year', 'month', 'created', 'updated'.
    Also detect if >70% of non-null values in a column match a common date pattern.
15. For detected date columns: use pd.to_datetime(df[col], infer_datetime_format=True, errors='coerce')
    then format to ISO 8601: df[col].dt.strftime('%Y-%m-%d').
    Leave failed parses as NaN — do NOT drop those rows.

── CATEGORICAL / STRING COLUMNS ──
16. Normalize casing per column:
    - Proper nouns (names, cities): .str.title()
    - Codes, IDs, categories: .str.upper()
    - General free text: .str.lower()
17. For low-cardinality columns (< 30 unique values after stripping):
    group near-duplicate categories caused by casing/spacing inconsistencies
    (e.g. 'Male', 'male', 'MALE', 'M') by mapping variants to the most frequent form.
    Use value_counts() to find the dominant form.

── BOOLEAN-LIKE COLUMNS ──
18. If a column contains values only from the set {{yes, no, true, false, 1, 0, y, n, t, f}}
    (case-insensitive), map them all to Python bool True/False.

── ID / CODE COLUMNS ──
19. If a column name contains 'id', 'code', 'zip', 'pin', 'sku', 'ref', or 'number':
    preserve as string, do NOT cast to numeric, do NOT strip leading zeros.
    Cast to str and strip whitespace only.

── CONTACT / URL COLUMNS ──
20. Phone columns (name contains 'phone', 'mobile', 'tel', 'fax'):
    strip all non-digit characters except a leading '+'. Keep as string.
21. Email columns (name contains 'email', 'mail'):
    apply .str.lower().str.strip(). Print a warning for rows where value lacks '@'.
22. URL columns (name contains 'url', 'link', 'website'):
    apply .str.lower().str.strip().

── POST-CLEAN VALIDATION ──
23. After all cleaning, assert the DataFrame is not empty:
    if df.empty: raise ValueError("Cleaning resulted in an empty DataFrame — aborting.")
24. At the end, print a concise cleaning report:
    print(f"Cleaning complete: {{original_rows}} rows → {{len(df)}} rows | {{original_cols}} cols → {{len(df.columns)}} cols")

STRICT RULES:
- ALWAYS check `if df[col].dtype == object` before applying .str methods.
- Use df[col].str methods for string operations — NEVER df.applymap() or df.map() on the whole DataFrame.
- Do NOT rename or reorder columns beyond snake_case conversion.
- Do NOT aggregate, pivot, merge, or reshape the data in any way.
- Do NOT drop rows for outliers — cap (Winsorize) them instead.
- imports required at the top of the function body: import pandas as pd, import numpy as np, import re.
- The function must accept a single argument `df` and return only the cleaned DataFrame.
- Return ONLY valid, executable Python code. NO markdown. NO backticks. NO explanations outside print().
"""

                    try:
                        clean_response = model.generate_content(clean_prompt)
                        clean_code = (
                            clean_response.text.strip()
                            .replace("```python", "")
                            .replace("```", "")
                            .strip()
                        )

                        local_vars = {}
                        exec(clean_code, globals(), local_vars)
                        df_cleaned = local_vars['clean_data'](df_to_process)

                        # ── Before / After ────────────────────────────
                        st.subheader(f"Results for: {file}")
                        b_col, a_col = st.columns(2)
                        with b_col:
                            st.markdown("**🔴 Before (Raw)**")
                            st.dataframe(df_to_process.head(5), use_container_width=True)
                            st.caption(f"{original_rows:,} rows · {original_cols} columns")
                        with a_col:
                            st.markdown("**🟢 After (Sanitised)**")
                            st.dataframe(df_cleaned.head(5), use_container_width=True)
                            st.caption(f"{len(df_cleaned):,} rows · {len(df_cleaned.columns)} columns")

                        # ── Auto-save to Supabase ─────────────────────
                        clean_name       = f"sanitized_{file}"
                        df_cleaned_json  = df_cleaned.replace({np.nan: None})
                        supabase.table('cleaned_datasets').insert({
                            'file_name':     clean_name,
                            'data':          df_cleaned_json.to_dict(orient='records'),
                            'user_id':       user_id,
                            'notebook_name': nb_name
                        }).execute()
                        st.session_state['clean_cloud'][clean_name] = df_cleaned
                        st.success(f"✅ '{file}' sanitised and auto-saved to database as '{clean_name}'.")

                        # ── Format-aware download ─────────────────────
                        file_ext_dl = pathlib.Path(file).suffix.lower()
                        try:
                            if file_ext_dl == '.csv':
                                file_bytes = df_cleaned.to_csv(index=False).encode('utf-8')
                                mime       = "text/csv"
                                dl_name    = clean_name
                            elif file_ext_dl in ['.xlsx', '.xls']:
                                buf = io.BytesIO()
                                with pd.ExcelWriter(buf, engine='openpyxl') as writer:
                                    df_cleaned.to_excel(writer, index=False)
                                file_bytes = buf.getvalue()
                                mime       = ("application/vnd.openxmlformats-officedocument"
                                              ".spreadsheetml.sheet")
                                dl_name    = f"sanitized_{pathlib.Path(file).stem}.xlsx"
                            elif file_ext_dl == '.json':
                                file_bytes = df_cleaned.to_json(orient='records', indent=2).encode('utf-8')
                                mime       = "application/json"
                                dl_name    = clean_name
                            else:
                                # Fallback for .txt, .pdf, .docx, .db, .sqlite
                                file_bytes = df_cleaned.to_csv(index=False).encode('utf-8')
                                mime       = "text/csv"
                                dl_name    = f"sanitized_{pathlib.Path(file).stem}.csv"
                        except ImportError:
                            file_bytes = df_cleaned.to_csv(index=False).encode('utf-8')
                            mime       = "text/csv"
                            dl_name    = f"sanitized_{pathlib.Path(file).stem}.csv"

                        st.download_button(
                            label=f"⬇️ Download {dl_name}",
                            data=file_bytes,
                            file_name=dl_name,
                            mime=mime,
                            use_container_width=True
                        )
                        st.divider()

                    except Exception as e:
                        st.error(f"🚨 Sanitization failed for '{file}': {e}")

                    progress_bar.progress((i + 1) / total_files)

                status_text.success("🎉 All selected files have been sanitised and saved!")

    # ── Chat Section ───────────────────────────────────────────
    st.divider()
    st.header("💬 Chat with your Data")

    all_chat_data = {
        **st.session_state.get('raw_cloud', {}),
        **st.session_state.get('clean_cloud', {})
    }

    if not all_chat_data:
        st.info("Upload or sanitise data above to begin chatting.")
    else:
        chat_file = st.selectbox(
            "Which dataset do you want to talk to?",
            list(all_chat_data.keys()),
            key="chat_file_selector"
        )
        df_chat = all_chat_data[chat_file]

        llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash",
            temperature=0,
            google_api_key=st.secrets["GEMINI_API_KEY"]
        )
        agent = create_pandas_dataframe_agent(
            llm, df_chat,
            verbose=True,
            allow_dangerous_code=True,
            handle_parsing_errors=True
        )

        if "messages" not in st.session_state:
            st.session_state.messages = []

        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        if prompt := st.chat_input(f"Ask something about {chat_file}..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            with st.chat_message("assistant"):
                with st.spinner("Crunching the numbers..."):
                    try:
                        response = agent.invoke(prompt)
                        answer   = response["output"]
                        st.markdown(answer)
                        st.session_state.messages.append({"role": "assistant", "content": answer})
                    except Exception as e:
                        error_str = str(e)
                        if "Could not parse LLM output:" in error_str:
                            extracted = (
                                error_str
                                .split("Could not parse LLM output:")[1]
                                .split("For troubleshooting, visit:")[0]
                                .strip(" `\n")
                            )
                            st.markdown(extracted)
                            st.session_state.messages.append({"role": "assistant", "content": extracted})
                        else:
                            st.error(f"Could not calculate that. Error: {e}")


# ══════════════════════════════════════════════════════════════
# TAB 3 — DASHBOARD
# ══════════════════════════════════════════════════════════════
with tab3:
    all_dash_data = {
        **st.session_state.get('raw_cloud', {}),
        **st.session_state.get('clean_cloud', {})
    }

    if not all_dash_data:
        st.info("⬆️ No data yet — upload your files in **Part 1** first.")
    else:
        st.header("📊 Multi-File Dashboard")

        dash_files = st.multiselect(
            "Select dataset(s) to include in the dashboard:",
            list(all_dash_data.keys()),
            default=list(all_dash_data.keys())[:1]
        )

        if dash_files:
            df_dict       = {f: all_dash_data[f] for f in dash_files}
            total_records = sum(len(df) for df in df_dict.values())
            total_features = sum(len(df.columns) for df in df_dict.values())

            st.subheader("Aggregated Data Scope")
            kpi1, kpi2 = st.columns(2)
            kpi1.metric("Total Records Across Files",  f"{total_records:,}")
            kpi2.metric("Total Features Across Files", total_features)
            st.divider()

            if st.button("🪄 Auto-Generate Full Dashboard", type="primary"):
                with st.spinner("The AI is analysing relationships across your files..."):
                    schema_context = "\n".join([
                        f"- {name}: columns={list(df.columns)}, rows={len(df)}"
                        for name, df in df_dict.items()
                    ])
                    sample_context = "\n\n".join([
                        f"### {name} (first 5 rows):\n{df.head(5).to_csv(index=False)}"
                        for name, df in df_dict.items()
                    ])

                    dash_prompt = f"""
You are a Senior BI Developer and Data Visualisation expert.
Write a Python function `build_dashboard(data_dict)` that takes a dictionary of pandas DataFrames
and generates exactly 4 insightful, distinct Plotly charts tailored to the actual data provided.

INPUT DATA SCHEMA:
{schema_context}

SAMPLE DATA:
{sample_context}

THE FUNCTION MUST:
1. Import plotly.express as px and plotly.graph_objects as go at the top of the function.
2. Extract DataFrames from data_dict using the exact key names listed above.
3. Inspect column names and dtypes to choose appropriate chart types automatically.
4. Create 4 DIFFERENT chart types (choose from: bar, line, scatter, pie, histogram,
   box, heatmap, treemap, funnel) — do not repeat the same type twice.
5. Each chart must have a descriptive title, axis labels, and use a clean colour palette.
6. If multiple DataFrames are present, at least one chart must combine or compare data
   across two or more files (e.g., merged on a common key column).
7. Handle potential errors gracefully: if a chart cannot be built due to incompatible data,
   create a simple placeholder bar chart of value counts for the first column instead.
8. Return exactly this structure:
   {{"Chart Title 1": fig1, "Chart Title 2": fig2, "Chart Title 3": fig3, "Chart Title 4": fig4}}

STRICT RULES:
- Do NOT use df.applymap() or any deprecated pandas methods.
- Do NOT hardcode column names — infer them from the DataFrames at runtime.
- Return ONLY valid executable Python code. NO markdown. NO backticks. NO explanations.
"""
                    try:
                        dash_response = model.generate_content(dash_prompt)
                        dash_code = (
                            dash_response.text.strip()
                            .replace("```python", "")
                            .replace("```", "")
                            .strip()
                        )

                        local_vars = {}
                        exec(dash_code, globals(), local_vars)
                        dashboard_figs = local_vars['build_dashboard'](df_dict)

                        st.success("✅ Dashboard Generated!")
                        chart_titles = list(dashboard_figs.keys())

                        row1_col1, row1_col2 = st.columns(2)
                        with row1_col1:
                            st.subheader(chart_titles[0])
                            st.plotly_chart(dashboard_figs[chart_titles[0]], use_container_width=True)
                        with row1_col2:
                            st.subheader(chart_titles[1])
                            st.plotly_chart(dashboard_figs[chart_titles[1]], use_container_width=True)

                        row2_col1, row2_col2 = st.columns(2)
                        with row2_col1:
                            st.subheader(chart_titles[2])
                            st.plotly_chart(dashboard_figs[chart_titles[2]], use_container_width=True)
                        with row2_col2:
                            st.subheader(chart_titles[3])
                            st.plotly_chart(dashboard_figs[chart_titles[3]], use_container_width=True)
'''

import streamlit as st
import pandas as pd
import pathlib
import sqlite3
import tempfile
import PyPDF2
import docx
import re
import numpy as np
import json
import io
import time
import plotly.express as px
import google.generativeai as genai
from supabase import create_client, Client
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_experimental.agents import create_pandas_dataframe_agent

st.set_page_config(page_title="AI Data Agent SaaS", layout="wide")

# ──────────────────────────────────────────────────────────────
# 1. SETUP THE BRAIN & THE VAULT
# ──────────────────────────────────────────────────────────────
try:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])
    model = genai.GenerativeModel('gemini-3-flash-preview')
    supabase: Client = create_client(st.secrets["SUPABASE_URL"], st.secrets["SUPABASE_KEY"])

    if 'session' in st.session_state and st.session_state['session'] is not None:
        try:
            supabase.postgrest.auth(st.session_state['session'].access_token)
            supabase.auth.set_session(
                st.session_state['session'].access_token,
                st.session_state['session'].refresh_token
            )
        except Exception:
            st.session_state['user'] = None
            st.session_state['session'] = None
except Exception as e:
    st.error(f"Setup Error: Please check your Streamlit Secrets. ({e})")
    st.stop()

# ──────────────────────────────────────────────────────────────
# 2. THE BOUNCER (LOGIN / REGISTER)
# ──────────────────────────────────────────────────────────────
if 'user' not in st.session_state:
    st.session_state['user'] = None
if 'session' not in st.session_state:
    st.session_state['session'] = None

if st.session_state['user'] is None:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.title("🔐 Welcome to AI Data Agent")
        st.markdown("Please log in to access your secure workspace.")

        auth_tab1, auth_tab2 = st.tabs(["Login", "Register"])
        with auth_tab1:
            log_email = st.text_input("Email", key="log_email")
            log_pass  = st.text_input("Password", type="password", key="log_pass")
            if st.button("Login", use_container_width=True, type="primary"):
                try:
                    response = supabase.auth.sign_in_with_password({"email": log_email, "password": log_pass})
                    st.session_state['user']    = response.user
                    st.session_state['session'] = response.session
                    st.rerun()
                except Exception:
                    st.error("Login failed. Please check your credentials.")

        with auth_tab2:
            reg_name  = st.text_input("First Name", key="reg_name")
            reg_email = st.text_input("New Email", key="reg_email")
            reg_pass  = st.text_input("New Password", type="password", key="reg_pass")
            if st.button("Register", use_container_width=True):
                try:
                    supabase.auth.sign_up({
                        "email": reg_email, "password": reg_pass,
                        "options": {"data": {"first_name": reg_name}}
                    })
                    st.success("Registration successful! You can now log in.")
                except Exception as e:
                    st.error(f"Registration failed: {e}")
    st.stop()

# ──────────────────────────────────────────────────────────────
# 2.5  SESSION RESTORE  (runs once per login / after any refresh)
# ──────────────────────────────────────────────────────────────
_uid = st.session_state['user'].id

if not st.session_state.get('_session_restored'):
    try:
        _raw_nb   = supabase.table("raw_datasets").select("notebook_name").eq("user_id", _uid).execute()
        _clean_nb = supabase.table("cleaned_datasets").select("notebook_name").eq("user_id", _uid).execute()
        _all_nb   = sorted(set(
            [r['notebook_name'] for r in _raw_nb.data] +
            [r['notebook_name'] for r in _clean_nb.data]
        ))
        if _all_nb and not st.session_state.get('notebook_list'):
            st.session_state['notebook_list'] = _all_nb
        if st.session_state.get('notebook_list') and st.session_state.get('active_notebook') is None:
            st.session_state['active_notebook'] = st.session_state['notebook_list'][0]
        st.session_state['_session_restored'] = True
    except Exception:
        st.session_state['_session_restored'] = True

# ──────────────────────────────────────────────────────────────
# 3. THE SIDEBAR (PROFILE & NOTEBOOKS)
# ──────────────────────────────────────────────────────────────
user_name = st.session_state['user'].user_metadata.get('first_name', 'User')

st.sidebar.markdown("### 👤 My Profile")
st.sidebar.markdown(f"**Name:** {user_name}")
st.sidebar.markdown(f"**Email:** {st.session_state['user'].email}")

if st.sidebar.button("🚪 Log Out", use_container_width=True):
    for _k in ['user', 'session', 'notebook_list', 'active_notebook',
               'raw_cloud', 'clean_cloud', 'messages', '_session_restored', '_loaded_nb']:
        st.session_state.pop(_k, None)
    supabase.auth.sign_out()
    st.rerun()

st.sidebar.divider()

if 'notebook_list'  not in st.session_state: st.session_state['notebook_list']  = []
if 'active_notebook' not in st.session_state: st.session_state['active_notebook'] = None
if 'show_nb_modal'  not in st.session_state: st.session_state['show_nb_modal']  = False
if 'confirm_del_nb' not in st.session_state: st.session_state['confirm_del_nb'] = False

st.sidebar.markdown("### 📚 My Notebooks")

# ── New Notebook Button ────────────────────────────────────────
if st.sidebar.button("📓 New Notebook", use_container_width=True):
    st.session_state['show_nb_modal']  = True
    st.session_state['confirm_del_nb'] = False

# ── New Notebook Modal ─────────────────────────────────────────
@st.dialog("Create a New Notebook")
def new_notebook_modal():
    nb_input = st.text_input("Notebook name", placeholder="e.g., Q3 Financials")
    col_ok, col_cancel = st.columns(2)
    with col_ok:
        if st.button("✅ Create", use_container_width=True, type="primary"):
            if nb_input and nb_input not in st.session_state['notebook_list']:
                st.session_state['notebook_list'].append(nb_input)
                st.session_state['active_notebook'] = nb_input
                st.session_state['show_nb_modal']   = False
                st.rerun()
            elif nb_input in st.session_state['notebook_list']:
                st.warning("A notebook with this name already exists.")
            else:
                st.warning("Please enter a name.")
    with col_cancel:
        if st.button("Cancel", use_container_width=True):
            st.session_state['show_nb_modal'] = False
            st.rerun()

if st.session_state['show_nb_modal']:
    new_notebook_modal()

# ── Notebook List: select button + 🗑 delete icon per row ─────
if st.session_state['notebook_list']:
    st.sidebar.divider()
    st.sidebar.markdown("**My Notebooks:**")

    for nb in st.session_state['notebook_list']:
        is_active  = (nb == st.session_state.get('active_notebook'))
        col_nm, col_dl = st.sidebar.columns([5, 1])
        with col_nm:
            label = f"**› {nb}**" if is_active else nb
            if st.button(label, key=f"nb_select_{nb}", use_container_width=True):
                st.session_state['active_notebook'] = nb
                st.session_state['confirm_del_nb']  = False
                st.rerun()
        with col_dl:
            if st.button("🗑", key=f"nb_del_{nb}", help=f"Delete notebook '{nb}'"):
                st.session_state['confirm_del_nb'] = True
                st.session_state['nb_to_delete']   = nb
                st.rerun()

# ── Delete Notebook Confirmation Dialog ───────────────────────
@st.dialog("⚠️ Confirm Notebook Deletion")
def confirm_delete_notebook():
    nb_del = st.session_state.get('nb_to_delete', '')
    st.warning(
        f"Are you sure you want to permanently delete **'{nb_del}'**?\n\n"
        f"This will remove **all raw and cleaned datasets** stored in this notebook. "
        f"This action **cannot be undone**."
    )
    col_yes, col_no = st.columns(2)
    with col_yes:
        if st.button("🗑️ Yes, Delete", use_container_width=True, type="primary"):
            try:
                supabase.table("raw_datasets").delete() \
                    .eq("user_id", _uid).eq("notebook_name", nb_del).execute()
                supabase.table("cleaned_datasets").delete() \
                    .eq("user_id", _uid).eq("notebook_name", nb_del).execute()

                st.session_state['notebook_list'].remove(nb_del)
                remaining = st.session_state['notebook_list']
                st.session_state['active_notebook'] = remaining[0] if remaining else None
                for _k in ['raw_cloud', 'clean_cloud', '_loaded_nb']:
                    st.session_state.pop(_k, None)
                st.session_state['confirm_del_nb'] = False
                st.session_state.pop('nb_to_delete', None)
                st.rerun()
            except Exception as e:
                st.error(f"Could not delete notebook: {e}")
    with col_no:
        if st.button("Cancel", use_container_width=True):
            st.session_state['confirm_del_nb'] = False
            st.session_state.pop('nb_to_delete', None)
            st.rerun()

if st.session_state.get('confirm_del_nb'):
    confirm_delete_notebook()

# ──────────────────────────────────────────────────────────────
# 4. THE NOTEBOOK ENFORCER
# ──────────────────────────────────────────────────────────────
if st.session_state['active_notebook'] is None:
    st.warning("👋 Welcome! Please create your first Notebook in the sidebar to begin working.")
    st.stop()

st.title(f"📂 Notebook: {st.session_state['active_notebook']}")
st.divider()

# ──────────────────────────────────────────────────────────────
# 5. SECURE CLOUD FETCHING (only re-fetches when notebook changes)
# ──────────────────────────────────────────────────────────────
user_id = st.session_state['user'].id
nb_name = st.session_state['active_notebook']

_nb_changed = (st.session_state.get('_loaded_nb') != nb_name)

if _nb_changed or 'raw_cloud' not in st.session_state:
    st.session_state['raw_cloud']   = {}
    st.session_state['clean_cloud'] = {}
    with st.spinner(f"Loading {nb_name} vault..."):
        try:
            raw_res = supabase.table("raw_datasets").select("file_name, data") \
                .eq("user_id", user_id).eq("notebook_name", nb_name).execute()
            st.session_state['raw_cloud'] = {
                r['file_name']: pd.DataFrame(r['data']) for r in raw_res.data
            }
            clean_res = supabase.table("cleaned_datasets").select("file_name, data") \
                .eq("user_id", user_id).eq("notebook_name", nb_name).execute()
            st.session_state['clean_cloud'] = {
                r['file_name']: pd.DataFrame(r['data']) for r in clean_res.data
            }
            st.session_state['_loaded_nb'] = nb_name
        except Exception as e:
            st.warning(f"Could not fetch cloud data: {e}")

# ──────────────────────────────────────────────────────────────
# TABBED ARCHITECTURE
# ──────────────────────────────────────────────────────────────
tab1, tab2, tab3 = st.tabs([
    "📤 Part 1 — Upload",
    "🧹 Part 2 — Sanitise & Query",
    "📊 Part 3 — Dashboard"
])


# ══════════════════════════════════════════════════════════════
# TAB 1 — UPLOAD
# ══════════════════════════════════════════════════════════════
with tab1:
    st.subheader("📤 Upload New Data")
    st.caption(
        "Supported: CSV, Excel (.xlsx/.xls), JSON, SQLite (.db/.sqlite), "
        "PDF, Word (.docx), plain text (.txt). Multiple files at once supported."
    )

    uploaded_files = st.file_uploader(
        f"Drop files into '{nb_name}'",
        type=["csv", "xlsx", "xls", "json", "db", "sqlite", "pdf", "docx", "txt"],
        accept_multiple_files=True
    )

    if st.button("🚀 Upload File(s) to Database", type="primary"):
        if not uploaded_files:
            st.warning("Please select at least one file to upload.")
        else:
            for file in uploaded_files:

                # ── GUARD 1: Already in this notebook ─────────────────
                if file.name in st.session_state.get('raw_cloud', {}):
                    st.warning(
                        f"⚠️ **'{file.name}'** already exists in this notebook. "
                        f"Skipped. To replace it, delete it in the file manager below."
                    )
                    continue

                # ── GUARD 2: Exists in another notebook — reuse ───────
                try:
                    existing_res = supabase.table("raw_datasets") \
                        .select("file_name, data, notebook_name") \
                        .eq("user_id", user_id).eq("file_name", file.name) \
                        .neq("notebook_name", nb_name).limit(1).execute()
                    if existing_res.data:
                        source_nb = existing_res.data[0]['notebook_name']
                        reused_df = pd.DataFrame(existing_res.data[0]['data']).replace({np.nan: None})
                        supabase.table('raw_datasets').insert({
                            'file_name': file.name, 'data': reused_df.to_dict(orient='records'),
                            'user_id': user_id, 'notebook_name': nb_name
                        }).execute()
                        st.session_state['raw_cloud'][file.name] = reused_df
                        st.info(f"♻️ **'{file.name}'** found in notebook **'{source_nb}'** — linked here.")
                        continue
                except Exception:
                    pass  # fall through to normal upload

                # ── NORMAL PARSE + UPLOAD ──────────────────────────────
                file_ext = pathlib.Path(file.name).suffix.lower()
                try:
                    if file_ext in ['.csv', '.xlsx', '.xls', '.json']:
                        if file_ext == '.csv':
                            df = pd.read_csv(file)
                        elif file_ext in ['.xlsx', '.xls']:
                            df = pd.read_excel(file)
                        elif file_ext == '.json':
                            try:
                                df = pd.read_json(file, orient='records')
                            except Exception:
                                file.seek(0)
                                try:
                                    df = pd.read_json(file, lines=True)
                                except Exception:
                                    file.seek(0)
                                    content = file.getvalue().decode("utf-8")
                                    objs = []
                                    for line in content.splitlines():
                                        if line.strip():
                                            try: objs.append(json.loads(line))
                                            except json.JSONDecodeError: pass
                                    if objs: df = pd.json_normalize(objs)
                                    else: raise Exception("Could not extract tabular JSON data.")
                        df = df.replace({np.nan: None})
                        with st.spinner(f"Saving {file.name}..."):
                            supabase.table('raw_datasets').insert({
                                'file_name': file.name, 'data': df.to_dict(orient='records'),
                                'user_id': user_id, 'notebook_name': nb_name
                            }).execute()
                        st.session_state['raw_cloud'][file.name] = df
                        st.success(f"✅ Uploaded: {file.name}")

                    elif file_ext in ['.txt', '.pdf', '.docx', '.db', '.sqlite']:
                        if file_ext == '.txt':
                            text_data = file.getvalue().decode("utf-8")
                        elif file_ext == '.pdf':
                            pdf_reader = PyPDF2.PdfReader(file)
                            text_data  = "".join([p.extract_text() + "\n" for p in pdf_reader.pages])
                        elif file_ext == '.docx':
                            doc_obj   = docx.Document(file)
                            text_data = "\n".join([p.text for p in doc_obj.paragraphs])
                        elif file_ext in ['.db', '.sqlite']:
                            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp:
                                tmp.write(file.read()); tmp_path = tmp.name
                            conn = sqlite3.connect(tmp_path)
                            tables = pd.read_sql("SELECT name FROM sqlite_master WHERE type='table'", conn)
                            text_data = f"SQLite DB: {file.name}\nTables: {tables['name'].tolist()}"
                            conn.close()
                        df = pd.DataFrame([{"document_name": file.name, "content": text_data}])
                        with st.spinner(f"Saving {file.name}..."):
                            supabase.table('raw_datasets').insert({
                                'file_name': file.name, 'data': df.to_dict(orient='records'),
                                'user_id': user_id, 'notebook_name': nb_name
                            }).execute()
                        st.session_state['raw_cloud'][file.name] = df
                        st.success(f"✅ Uploaded document: {file.name}")

                except Exception as e:
                    st.error(f"Error uploading {file.name}: {e}")

    # ── Success nudge ──────────────────────────────────────────
    if st.session_state.get('raw_cloud'):
        st.success(
            f"✅ {len(st.session_state['raw_cloud'])} file(s) ready — "
            f"head to **Part 2** to sanitise or **Part 3** for your dashboard."
        )

    # ── File Previews ──────────────────────────────────────────
    if st.session_state.get('raw_cloud'):
        st.divider()
        st.subheader("🗃️ Uploaded Files Preview")
        for fname, fdf in st.session_state['raw_cloud'].items():
            with st.expander(f"📄 {fname}  —  {len(fdf):,} rows × {len(fdf.columns)} cols"):
                st.dataframe(fdf.head(5), use_container_width=True)

    # ══ FILE MANAGER — Checkbox-based multi-delete ════════════
    st.divider()
    st.subheader("🗂️ Manage Files")

    all_stored_files = (
        [("raw_datasets",     k, "🟡 RAW")   for k in st.session_state.get('raw_cloud',   {}).keys()] +
        [("cleaned_datasets", k, "🟢 CLEAN") for k in st.session_state.get('clean_cloud', {}).keys()]
    )

    if not all_stored_files:
        st.caption("No files in this notebook yet.")
    else:
        # Select All toggle
        select_all = st.checkbox("☑️ Select All", key="del_select_all")

        checked_files = []
        for table_name, fname, badge in all_stored_files:
            col_chk, col_label = st.columns([1, 11])
            with col_chk:
                checked = st.checkbox(
                    "", key=f"del_chk_{table_name}_{fname}",
                    value=select_all, label_visibility="collapsed"
                )
            with col_label:
                st.markdown(f"{badge} &nbsp; `{fname}`", unsafe_allow_html=True)
            if checked:
                checked_files.append((table_name, fname))

        st.write("")
        if checked_files:
            st.caption(f"**{len(checked_files)}** file(s) selected for deletion")
            if st.button(
                f"🗑️ Delete {len(checked_files)} Selected File(s)",
                type="secondary", use_container_width=True
            ):
                errors = []
                for table_name, original_name in checked_files:
                    try:
                        supabase.table(table_name).delete() \
                            .eq("user_id",       user_id) \
                            .eq("notebook_name", nb_name) \
                            .eq("file_name",     original_name) \
                            .execute()
                        if table_name == "raw_datasets":
                            st.session_state['raw_cloud'].pop(original_name, None)
                        else:
                            st.session_state['clean_cloud'].pop(original_name, None)
                    except Exception as e:
                        errors.append(f"{original_name}: {e}")

                if errors:
                    for err in errors: st.error(f"🚨 {err}")
                else:
                    st.success(f"✅ {len(checked_files)} file(s) permanently deleted.")
                st.rerun()


# ══════════════════════════════════════════════════════════════
# TAB 2 — SANITISE & CHAT
# ══════════════════════════════════════════════════════════════
with tab2:

    if not st.session_state.get('raw_cloud'):
        st.info("⬆️ No data yet — upload your files in **Part 1** first.")

    # ── Sanitizer — sequential multi-file loop ─────────────────
    st.header("🧠 Universal Data Sanitizer")

    raw_keys = list(st.session_state.get('raw_cloud', {}).keys())
    if raw_keys:
        selected_files = st.multiselect(
            "Select raw dataset(s) to sanitise — each file processed in sequence:",
            raw_keys,
            default=raw_keys[:1]
        )

        if st.button("🪄 Run Exhaustive AI Sanitization", type="primary"):
            if not selected_files:
                st.warning("Please select at least one file to sanitize.")
            else:
                total_files  = len(selected_files)
                progress_bar = st.progress(0)
                status_text  = st.empty()

                for i, file in enumerate(selected_files):
                    est_time = (total_files - i) * 15
                    status_text.info(
                        f"Processing **{i+1} of {total_files}**: `{file}` "
                        f"— Est. remaining: ~{est_time}s"
                    )

                    df_to_process = st.session_state['raw_cloud'][file]
                    original_rows = len(df_to_process)
                    original_cols = len(df_to_process.columns)

                    # ── Rich column profile for the AI ────────
                    profile_lines = []
                    for col in df_to_process.columns:
                        dtype       = str(df_to_process[col].dtype)
                        null_pct    = round(df_to_process[col].isnull().mean() * 100, 1)
                        sample_vals = df_to_process[col].dropna().astype(str).unique()[:5].tolist()
                        profile_lines.append(
                            f"  - {col} | dtype={dtype} | nulls={null_pct}% | samples={sample_vals}"
                        )
                    col_profile = "\n".join(profile_lines)
                    full_sample = df_to_process.head(20).to_csv(index=False)

                    clean_prompt = f"""
You are a Senior Data Engineer and experienced Data Analyst. Write a single Python function
called `clean_data(df)` that exhaustively sanitizes the provided dataset for analysis.

DATASET PROFILE (use this to make targeted, column-level decisions):
{col_profile}

SAMPLE DATA (first 20 rows):
{full_sample}

YOUR FUNCTION MUST HANDLE ALL OF THE FOLLOWING — apply only what is relevant to this data:

── STRUCTURAL ──
1. Convert all column headers to strict snake_case: lowercase, strip whitespace,
   replace spaces/special chars with underscores, collapse multiple underscores.
2. Drop columns that are 100% empty (all NaN/None).
3. Drop rows that are 100% empty.
4. Drop exact duplicate rows.
5. Reset the DataFrame index after dropping.

── WHITESPACE & ENCODING ──
6. For every object/string column: strip leading/trailing whitespace via df[col].str.strip().
7. Remove hidden chars: null bytes (\\x00), BOM (\\ufeff), zero-width spaces (\\u200b),
   carriage returns (\\r). Use df[col].str.replace() with a compiled regex per column.
8. Normalize multiple consecutive internal spaces to single space.

── NULL / MISSING VALUES ──
9. Replace junk null strings with np.nan (case-insensitive):
   '', 'n/a', 'na', 'none', 'null', 'nil', 'nan', '-', '--', '---', '?',
   'unknown', 'undefined', 'missing', 'not available', 'not applicable', '#n/a', '#null!'.
10. Numeric columns: null rate <20% impute median; 20-50% leave NaN;
    >50% drop column and print(f"Dropped high-null column: {{col}}").

── NUMERIC COLUMNS ──
11. Strip currency symbols ($,€,£,¥), comma separators, % from numeric-looking string
    columns and cast to float.
12. Object columns where all non-null values are numeric strings cast to float.
13. Outlier capping (Winsorize): for numeric cols where outlier count >1% of rows,
    cap at Q1-1.5*IQR and Q3+1.5*IQR boundaries. Never drop rows.

── DATE / TIME COLUMNS ──
14. Detect date cols by name keywords: 'date','time','dt','year','month','created','updated'.
    Also detect if >70% non-null values match a common date pattern.
15. pd.to_datetime(df[col], infer_datetime_format=True, errors='coerce'), then
    format ISO 8601: df[col].dt.strftime('%Y-%m-%d'). Leave parse failures as NaN.

── CATEGORICAL / STRING COLUMNS ──
16. Casing: names/cities -> .str.title() | codes/IDs -> .str.upper() | text -> .str.lower()
17. Low-cardinality cols (<30 unique): map near-duplicate variants to most-frequent form.

── BOOLEAN-LIKE COLUMNS ──
18. Cols with values only from {{yes,no,true,false,1,0,y,n,t,f}} (case-insensitive)
    map to Python bool True/False.

── ID / CODE COLUMNS ──
19. Cols whose name contains 'id','code','zip','pin','sku','ref','number':
    keep as string, no numeric cast, no leading-zero stripping.

── CONTACT / URL COLUMNS ──
20. Phone ('phone','mobile','tel','fax'): strip non-digit chars except leading '+'.
21. Email ('email','mail'): .str.lower().str.strip(); warn for rows missing '@'.
22. URL ('url','link','website'): .str.lower().str.strip().

── POST-CLEAN VALIDATION ──
23. if df.empty: raise ValueError("Cleaning resulted in empty DataFrame — aborting.")
24. print(f"Cleaning complete: {{original_rows}} -> {{len(df)}} rows | {{original_cols}} -> {{len(df.columns)}} cols")

STRICT RULES:
- Always check if df[col].dtype == object before applying .str methods.
- Use df[col].str methods — NEVER df.applymap() or df.map() on whole DataFrame.
- Do NOT aggregate, pivot, merge, or reshape.
- Imports inside function body: import pandas as pd, import numpy as np, import re.
- Return ONLY valid executable Python. NO markdown. NO backticks. NO explanations.
"""

                    try:
                        clean_response = model.generate_content(clean_prompt)
                        clean_code = (
                            clean_response.text.strip()
                            .replace("```python", "").replace("```", "").strip()
                        )
                        local_vars = {}
                        exec(clean_code, globals(), local_vars)
                        df_cleaned = local_vars['clean_data'](df_to_process)

                        # ── Before / After ────────────────────
                        st.subheader(f"Results for: `{file}`")
                        b_col, a_col = st.columns(2)
                        with b_col:
                            st.markdown("**🔴 Before (Raw)**")
                            st.dataframe(df_to_process.head(5), use_container_width=True)
                            st.caption(f"{original_rows:,} rows · {original_cols} cols")
                        with a_col:
                            st.markdown("**🟢 After (Sanitised)**")
                            st.dataframe(df_cleaned.head(5), use_container_width=True)
                            st.caption(f"{len(df_cleaned):,} rows · {len(df_cleaned.columns)} cols")

                        # ── Auto-save ─────────────────────────
                        clean_name      = f"sanitized_{file}"
                        df_cleaned_json = df_cleaned.replace({np.nan: None})
                        supabase.table('cleaned_datasets').insert({
                            'file_name': clean_name,
                            'data':      df_cleaned_json.to_dict(orient='records'),
                            'user_id':   user_id, 'notebook_name': nb_name
                        }).execute()
                        st.session_state['clean_cloud'][clean_name] = df_cleaned
                        st.success(f"✅ Auto-saved as **'{clean_name}'**.")

                        # ── Format-aware download ─────────────
                        ext_dl = pathlib.Path(file).suffix.lower()
                        try:
                            if ext_dl == '.csv':
                                fb = df_cleaned.to_csv(index=False).encode('utf-8')
                                mime = "text/csv"; dl_name = clean_name
                            elif ext_dl in ['.xlsx', '.xls']:
                                buf = io.BytesIO()
                                with pd.ExcelWriter(buf, engine='openpyxl') as w:
                                    df_cleaned.to_excel(w, index=False)
                                fb = buf.getvalue()
                                mime = ("application/vnd.openxmlformats-officedocument"
                                        ".spreadsheetml.sheet")
                                dl_name = f"sanitized_{pathlib.Path(file).stem}.xlsx"
                            elif ext_dl == '.json':
                                fb = df_cleaned.to_json(orient='records', indent=2).encode('utf-8')
                                mime = "application/json"; dl_name = clean_name
                            else:
                                fb = df_cleaned.to_csv(index=False).encode('utf-8')
                                mime = "text/csv"
                                dl_name = f"sanitized_{pathlib.Path(file).stem}.csv"
                        except ImportError:
                            fb = df_cleaned.to_csv(index=False).encode('utf-8')
                            mime = "text/csv"
                            dl_name = f"sanitized_{pathlib.Path(file).stem}.csv"

                        st.download_button(
                            label=f"⬇️ Download {dl_name}",
                            data=fb, file_name=dl_name, mime=mime,
                            use_container_width=True
                        )
                        st.divider()

                    except Exception as e:
                        st.error(f"🚨 Sanitization failed for '{file}': {e}")

                    progress_bar.progress((i + 1) / total_files)

                status_text.success(f"🎉 All {total_files} file(s) sanitised and saved!")

    # ── Chat Section ───────────────────────────────────────────
    st.divider()
    st.header("💬 Chat with your Data")

    all_chat_data = {
        **st.session_state.get('raw_cloud',   {}),
        **st.session_state.get('clean_cloud', {})
    }

    if not all_chat_data:
        st.info("Upload or sanitise data above to begin chatting.")
    else:
        chat_file = st.selectbox(
            "Which dataset do you want to talk to?",
            list(all_chat_data.keys()),
            key="chat_file_selector"
        )
        df_chat = all_chat_data[chat_file]

        llm = ChatGoogleGenerativeAI(
            model="gemini-2.0-flash", temperature=0,
            google_api_key=st.secrets["GEMINI_API_KEY"]
        )
        agent = create_pandas_dataframe_agent(
            llm, df_chat, verbose=True,
            allow_dangerous_code=True, handle_parsing_errors=True
        )

        if "messages" not in st.session_state:
            st.session_state.messages = []

        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        if prompt := st.chat_input(f"Ask something about {chat_file}..."):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            with st.chat_message("assistant"):
                with st.spinner("Crunching the numbers..."):
                    try:
                        response = agent.invoke(prompt)
                        answer   = response["output"]
                        st.markdown(answer)
                        st.session_state.messages.append({"role": "assistant", "content": answer})
                    except Exception as e:
                        error_str = str(e)
                        if "Could not parse LLM output:" in error_str:
                            extracted = (
                                error_str
                                .split("Could not parse LLM output:")[1]
                                .split("For troubleshooting, visit:")[0]
                                .strip(" `\n")
                            )
                            st.markdown(extracted)
                            st.session_state.messages.append({"role": "assistant", "content": extracted})
                        else:
                            st.error(f"Could not calculate that. Error: {e}")


# ══════════════════════════════════════════════════════════════
# TAB 3 — DASHBOARD
# ══════════════════════════════════════════════════════════════
with tab3:
    all_dash_data = {
        **st.session_state.get('raw_cloud',   {}),
        **st.session_state.get('clean_cloud', {})
    }

    if not all_dash_data:
        st.info("⬆️ No data yet — upload your files in **Part 1** first.")
    else:
        st.header("📊 Multi-File Dashboard")
        st.caption(
            "Select one or more datasets. Cleaned datasets recommended. "
            "The AI analyses column types, distributions, and cross-file relationships "
            "to generate a professional 6-chart dashboard."
        )

        # Default to cleaned files if available, else first raw file
        clean_keys   = list(st.session_state.get('clean_cloud', {}).keys())
        default_keys = clean_keys[:2] if clean_keys else list(all_dash_data.keys())[:1]

        # ── Multi-select: supports any number of files ─────────
        dash_files = st.multiselect(
            "Select dataset(s) to include in the dashboard:",
            list(all_dash_data.keys()),
            default=default_keys
        )

        if dash_files:
            df_dict        = {f: all_dash_data[f] for f in dash_files}
            total_records  = sum(len(df) for df in df_dict.values())
            total_features = sum(len(df.columns) for df in df_dict.values())

            # ── KPI row ────────────────────────────────────────
            k1, k2, k3 = st.columns(3)
            k1.metric("Datasets Selected",           len(dash_files))
            k2.metric("Total Records Across Files",  f"{total_records:,}")
            k3.metric("Total Features Across Files", total_features)
            st.divider()

            # ── Deep schema intelligence builder ───────────────
            def _build_schema_intel(df_dict):
                lines = []
                for name, df in df_dict.items():
                    lines.append(f"\n=== DATASET: {name} ===")
                    lines.append(f"Shape: {df.shape[0]} rows x {df.shape[1]} columns")
                    lines.append("Column details:")
                    for col in df.columns:
                        dtype    = str(df[col].dtype)
                        null_pct = round(df[col].isnull().mean() * 100, 1)
                        n_unique = df[col].nunique()
                        if pd.api.types.is_numeric_dtype(df[col]):
                            desc = df[col].describe()
                            lines.append(
                                f"  • {col} [NUMERIC] dtype={dtype} | nulls={null_pct}% | "
                                f"unique={n_unique} | min={round(desc['min'],2)} "
                                f"mean={round(desc['mean'],2)} max={round(desc['max'],2)} "
                                f"std={round(desc['std'],2)}"
                            )
                        elif pd.api.types.is_datetime64_any_dtype(df[col]):
                            lines.append(
                                f"  • {col} [DATETIME] | nulls={null_pct}% | "
                                f"range={df[col].min()} to {df[col].max()}"
                            )
                        else:
                            top_vals = df[col].value_counts().head(5).to_dict()
                            lines.append(
                                f"  • {col} [CATEGORICAL] dtype={dtype} | "
                                f"nulls={null_pct}% | unique={n_unique} | top={top_vals}"
                            )
                    lines.append(f"Sample (5 rows):\n{df.head(5).to_csv(index=False)}")
                return "\n".join(lines)

            def _find_join_candidates(df_dict):
                if len(df_dict) < 2: return []
                sets = [set(df.columns) for df in df_dict.values()]
                common = sets[0]
                for s in sets[1:]: common = common.intersection(s)
                return list(common)

            schema_intel    = _build_schema_intel(df_dict)
            join_candidates = _find_join_candidates(df_dict)
            join_hint = (
                f"Common columns that can serve as join keys: {join_candidates}"
                if join_candidates
                else "No common columns found — treat each dataset independently."
            )

            if st.button("🪄 Generate Professional Dashboard", type="primary"):
                with st.spinner("AI is analysing your data and building the dashboard..."):

                    dash_prompt = f"""
You are an expert Senior BI Developer, Data Scientist, and Data Visualisation specialist
with 15 years of experience building executive-grade dashboards in Plotly.

Your task: write a Python function `build_dashboard(data_dict)` that takes a dictionary
of pandas DataFrames and returns exactly 6 professional Plotly figures.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
DATASET INTELLIGENCE (read carefully before choosing charts):
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
{schema_intel}

JOIN HINT: {join_hint}

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
CHART SELECTION — choose based on column types in the data above:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- NUMERIC vs NUMERIC → scatter with trendline (px.scatter, trendline='ols')
- CATEGORICAL vs NUMERIC → horizontal sorted bar (px.bar, orientation='h')
- DATETIME vs NUMERIC → line chart showing trend (px.line, markers=True)
- SINGLE CATEGORICAL distribution → pie if <10 unique vals, else bar of value_counts
- NUMERIC distribution → histogram with box (px.histogram, marginal='box', nbins=30)
- MULTIPLE NUMERIC COLS → correlation heatmap (go.Heatmap on df.corr().round(2))
- MULTIPLE CATEGORIES vs NUMERIC → grouped bar (px.bar, barmode='group')
- If join possible → merge datasets on common key, then show a comparison chart

Pick the 6 most insightful chart types for THIS specific data.
Do NOT repeat the same Plotly function more than twice across 6 charts.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
QUALITY REQUIREMENTS — every chart must have all of these:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
1. Insight-driven title that describes the finding (NOT "Chart 1" or "Bar Chart")
   Examples: "Revenue by Region — Top 10", "Customer Age Distribution",
             "Sales vs Marketing Spend Correlation", "Monthly Signups Trend 2023"
2. Labelled axes with units where applicable (e.g. "Revenue (USD)", "Month", "Count")
3. Color palette: use px.colors.qualitative.Bold for categorical,
   px.colors.sequential.Blues for sequential data
4. Rich hover tooltips: set hover_data to include all relevant columns
5. Apply update_layout() to EVERY figure with these settings:
   - font=dict(family="Inter, sans-serif", size=12, color="#1a1a2e")
   - plot_bgcolor="#f8f9fa", paper_bgcolor="white"
   - margin=dict(l=50, r=30, t=70, b=50)
   - title_font=dict(size=15, color="#1a1a2e")
   - legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1)
   - showlegend=True
6. Bar charts: add text labels using text_auto='.2s', textposition='outside'
7. Line charts: set markers=True, line=dict(width=2)
8. Scatter plots: if a third numeric column exists, size=third_col_name
9. Heatmap: annotate cells with values, use colorscale='RdBu_r', zmid=0
10. Add a reference line to distributions: fig.add_vline(x=mean_val, line_dash='dash',
    annotation_text=f'Mean: {{mean_val:.1f}}', line_color='red')

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
FUNCTION CONTRACT:
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
- def build_dashboard(data_dict):
- Extract each df: df = data_dict["exact_key_from_schema_above"]
- ALWAYS infer column names at runtime — NEVER hardcode them:
  numeric_cols  = df.select_dtypes(include='number').columns.tolist()
  cat_cols      = df.select_dtypes(include='object').columns.tolist()
  datetime_cols = df.select_dtypes(include='datetime').columns.tolist()
  Then use numeric_cols[0], cat_cols[0], etc.
- Before every chart: df_plot = df.dropna(subset=[x_col, y_col]).copy()
  (or just df.dropna() for heatmaps/distributions)
- Wrap EACH chart block in its own try/except:
  except Exception:
      fig = px.bar(df[cat_cols[0]].value_counts().reset_index(),
                   x='count', y=cat_cols[0], orientation='h',
                   title=f"Data Overview: {{name}}")
- Return exactly this structure (use real descriptive titles as keys):
  {{
    "Descriptive Title 1": fig1,
    "Descriptive Title 2": fig2,
    "Descriptive Title 3": fig3,
    "Descriptive Title 4": fig4,
    "Descriptive Title 5": fig5,
    "Descriptive Title 6": fig6
  }}

STRICT CODE RULES:
- Imports at top of function: import plotly.express as px; import plotly.graph_objects as go
- Do NOT import pandas or numpy (already in scope as pd and np).
- Do NOT use df.applymap() — use df.map() or df[col].apply() instead.
- Do NOT hardcode any column names — always infer from dtypes at runtime.
- Return ONLY valid executable Python code. NO markdown. NO backticks. NO prose.
"""

                    try:
                        dash_response = model.generate_content(dash_prompt)
                        dash_code = (
                            dash_response.text.strip()
                            .replace("```python", "").replace("```", "").strip()
                        )
                        local_vars = {}
                        exec(dash_code, globals(), local_vars)
                        dashboard_figs = local_vars['build_dashboard'](df_dict)

                        n_charts     = len(dashboard_figs)
                        chart_titles = list(dashboard_figs.keys())
                        st.success(f"✅ Dashboard generated — {n_charts} charts ready!")

                        # Render in a 2-column grid, handles any number of charts
                        for row_start in range(0, n_charts, 2):
                            left_col, right_col = st.columns(2)
                            with left_col:
                                t = chart_titles[row_start]
                                st.subheader(t)
                                st.plotly_chart(dashboard_figs[t], use_container_width=True)
                            if row_start + 1 < n_charts:
                                with right_col:
                                    t = chart_titles[row_start + 1]
                                    st.subheader(t)
                                    st.plotly_chart(dashboard_figs[t], use_container_width=True)

                    except Exception as e:
                        st.error(f"🚨 Dashboard generation failed. Details: {e}")
                    except Exception as e:
                        st.error(f"🚨 The AI struggled to build the dashboard. Details: {e}")
